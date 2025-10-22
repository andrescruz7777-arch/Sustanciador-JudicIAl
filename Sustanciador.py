# ============================================================
# ⚖️ COS JudicIA — Sustanciador Judicial + Demandas & Medidas
# Autor: Andrés Cruz
# Fecha: 2025-10
# ============================================================

import io
import re
import unicodedata
from datetime import datetime
import zipfile
import pandas as pd
import streamlit as st
from docx import Document

# ==========================
# CONFIGURACIÓN GENERAL
# ==========================
st.set_page_config(page_title="⚖️ Sustanciador Judicial — COS JudicIA", layout="wide")
st.title("⚖️ COS JudicIA — Módulos Judiciales Automatizados")

tab1, tab2 = st.tabs(["⚙️ Sustanciador Judicial", "📄 Demandas y Medidas Cautelares"])

# ============================================================
# 🧩 TAB 1 — Sustanciador Judicial (actualizado)
# ============================================================
with tab1:
    st.header("⚙️ Sustanciador Judicial")
    st.caption("Genera memoriales desde plantillas .docx basadas en datos de Excel.")

    # ---------- Funciones auxiliares ----------
    def strip_accents(s: str) -> str:
        if not isinstance(s, str):
            return s
        return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

    def norm_text(s: str) -> str:
        return strip_accents(s or "").lower()

    def parse_ddmmyyyy(s: str):
        """Convierte fechas tipo dd/mm/yyyy o similares en objeto datetime."""
        if not isinstance(s, str):
            return None
        m = re.search(r"\b(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})\b", s)
        if not m:
            return None
        dd, mm, yyyy = m.groups()
        if len(yyyy) == 2:
            yyyy = f"20{yyyy}"
        try:
            return datetime(int(yyyy), int(mm), int(dd))
        except Exception:
            return None

    def format_fecha_dd_de_mm_de_yyyy(dt: datetime) -> str:
        if not dt:
            return ""
        meses = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        return f"{dt.day:02d} de {meses[dt.month-1]} de {dt.year}"

    def sanitize_filename(s: str) -> str:
        s = strip_accents(s or "").strip()
        s = re.sub(r"\s+", "_", s)
        s = re.sub(r"[^A-Za-z0-9._-]", "", s)
        return s

    def match_any(text: str, keys) -> bool:
        t = norm_text(text)
        return any(k in t for k in keys)

    def extract_fecha_mas_reciente_AF(series_like, keys):
        if series_like is None:
            return None
        candidatos = []
        if isinstance(series_like, (list, tuple, pd.Series)):
            textos = [str(x) for x in series_like if pd.notna(x)]
        else:
            textos = [str(series_like)] if pd.notna(series_like) else []
        for t in textos:
            if match_any(t, keys):
                dt = parse_ddmmyyyy(t)
                if dt:
                    candidatos.append(dt)
        return max(candidatos) if candidatos else None

    def replace_line_contains(doc: Document, contains_text: str, new_text: str):
        for p in doc.paragraphs:
            if contains_text in p.text:
                p.text = new_text
                return True
        return False

    def replace_after_label(doc: Document, label: str, new_value: str):
        label = label.rstrip(":")
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.upper().startswith(label.upper() + ":"):
                p.text = f"{label.upper()}: {new_value}"
                return True
        return False

    def replace_date_pattern(doc: Document, anchor_contains: str, pattern_regex: str, new_date_str: str):
        rx = re.compile(pattern_regex, flags=re.IGNORECASE)
        for p in doc.paragraphs:
            if anchor_contains.lower() in p.text.lower():
                new_text = rx.sub(new_date_str, p.text)
                if new_text != p.text:
                    p.text = new_text
                    return True
        return False

    def doc_to_preview_text(doc: Document) -> str:
        return "\n\n".join(p.text for p in doc.paragraphs)

    # ---------- Configuración de plantillas ----------
    MANDAMIENTO_KEYS = [
        "mandamiento", "correccion mandamiento", "correcion mandamiento",
        "solicitud correccion mandamiento", "solicitud correccion del mandamiento",
        "solicitud correccion del auto", "correccion del auto", "correcion del auto"
    ]

    SENTENCIA_KEYS = [
        "sentencia", "correccion sentencia", "correcion sentencia",
        "solicitud correccion sentencia", "solicitud correccion de sentencia",
        "solicitud correccion de la sentencia", "correccion del auto sentencia",
        "correcion del auto sentencia"
    ]

    TEMPLATES = {
        "Mandamiento": "MODELO IMPULSO PROCESAL CORRECCIÓN MP.docx",
        "Sentencia": "MODELO IMPULSO PROCESAL CORRECCIÓN SENTENCIA.docx",
        "Calificacion": "MODELO IMPULSO CALIFICACION DE DEMANDA.docx",
        "Liquidacion de credito": "MODELO IMPULSO LIQUIDACION DE CREDITO.docx",
    }

    SUBETAPAS = list(TEMPLATES.keys())

    PATTERN_PASADO = r"el pasado\s+\S+\s+de\s+\S+\s+de\s+\S+"
    PATTERN_RADICADA = r"radicada\s+el\s+d[ií]a\s+\S+\s+de\s+\S+\s+de\s+\S+"

    # ---------- UI Sustanciador ----------
    st.subheader("1️⃣ Carga la base en Excel")
    excel_file = st.file_uploader("Sube el archivo Excel con la precarga", type=["xlsx", "xls"])

    if excel_file:
        df = pd.read_excel(excel_file)
        st.success(f"Base cargada: {df.shape[0]} filas")
        st.dataframe(df.head(5))

        aliases = {
            'A': ['CC', 'Cédula', 'Documento'],
            'B': ['Nombre', 'Demandado'],
            'H': ['Juzgado'],
            'J': ['Correo', 'CorreoJuzgado'],
            'I': ['Radicado'],
            'O': ['FECHA DE PRESENTACIÓN DDA', 'FechaPresentacionDDA'],
            'AF': ['Cuaderno Principal']
        }

        col_map = {}
        lower_cols = {norm_text(c): c for c in df.columns}
        for key, names in aliases.items():
            for n in names:
                if norm_text(n) in lower_cols:
                    col_map[key] = lower_cols[norm_text(n)]
                    break

        subetapa = st.selectbox("Subetapa", SUBETAPAS)
        opciones = df.index.tolist()
        sel_idx = st.selectbox("Registro", opciones, format_func=lambda i: f"{df.loc[i, col_map.get('B')]} - {df.loc[i, col_map.get('I')]}")

        def generar_doc(row):
            cc = str(row.get(col_map.get('A'), ''))
            nombre = str(row.get(col_map.get('B'), ''))
            juzgado = str(row.get(col_map.get('H'), ''))
            correo = str(row.get(col_map.get('J'), ''))
            radicado = str(row.get(col_map.get('I'), ''))

            doc = Document(TEMPLATES[subetapa])
            replace_line_contains(doc, "JUZGADO", juzgado)
            replace_line_contains(doc, "@", correo)
            replace_after_label(doc, "RAD", radicado)
            replace_after_label(doc, "DEMANDANTE", "BANCO GNB SUDAMERIS S.A")
            replace_after_label(doc, "DEMANDADO", f"CC {cc} {nombre}")

            af_col = col_map.get('AF')
            o_col = col_map.get('O')

            # ====== Reemplazos por subetapa ======
            if subetapa == "Mandamiento" and af_col:
                fecha_dt = extract_fecha_mas_reciente_AF(row.get(af_col), MANDAMIENTO_KEYS)
                if fecha_dt:
                    replace_date_pattern(doc, "el pasado", PATTERN_PASADO,
                        f"el pasado {format_fecha_dd_de_mm_de_yyyy(fecha_dt)}")

            elif subetapa == "Sentencia" and af_col:
                fecha_dt = extract_fecha_mas_reciente_AF(row.get(af_col), SENTENCIA_KEYS)
                if fecha_dt:
                    replace_date_pattern(doc, "el pasado", PATTERN_PASADO,
                        f"el pasado {format_fecha_dd_de_mm_de_yyyy(fecha_dt)}")

            elif subetapa == "Calificacion" and o_col:
                raw_fecha = row.get(o_col)
                if isinstance(raw_fecha, datetime):
                    fecha_dt = raw_fecha
                else:
                    fecha_dt = parse_ddmmyyyy(str(raw_fecha))
                if fecha_dt:
                    fecha_str = format_fecha_dd_de_mm_de_yyyy(fecha_dt)
                    replace_date_pattern(doc, "radicada el", PATTERN_RADICADA,
                        f"radicada el día {fecha_str}")
            elif subetapa == "Liquidacion" and af_col:
                fecha_dt = extract_fecha_mas_reciente_AF(
                    row.get(af_col),
                    ["liquidacion", "liquidación", "credito", "crédito"
                )
                if fecha_dt:
                    fecha_str = format_fecha_dd_de_mm_de_yyyy(fecha_dt)
                    # Patrón flexible: acepta fecha real o los marcadores “xx de xxxx de xxxx”
                    pattern_liquidacion = (
                        r"(?i)(radicado\s+el\s+(?:d[ií]a\s+)?)(?:[0-9]{1,2}\s+de\s+\w+\s+de\s+[0-9]{4}|xx\s+de\s+\w+\s+de\s+xxxx)"
                    )
                    replace_date_pattern(
                        doc,
                        "radicado el",
                        pattern_liquidacion,
                        f"radicado el día {fecha_str}"
                    )
            
            return doc, cc, nombre

        # === Generación individual ===
        row = df.loc[sel_idx]
        doc, cc, nombre = generar_doc(row)
        st.subheader("Vista previa")
        st.text_area("Contenido", doc_to_preview_text(doc), height=300)

        out_name = f"{sanitize_filename(cc)}_{sanitize_filename(nombre)}_{subetapa}.docx"
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        st.download_button("⬇️ Descargar documento", data=bio, file_name=out_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # === Generación masiva ===
        st.divider()
        st.markdown("### 📦 Generación masiva (ZIP)")
        st.write("Genera todos los documentos de la base en un archivo ZIP.")

        if st.button("📦 Generar TODOS (ZIP)"):
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                for idx, fila in df.iterrows():
                    try:
                        doc, cc, nombre = generar_doc(fila)
                        out_name = f"{sanitize_filename(cc)}_{sanitize_filename(nombre)}_{subetapa}.docx"
                        out_mem = io.BytesIO()
                        doc.save(out_mem)
                        out_mem.seek(0)
                        zf.writestr(out_name, out_mem.read())
                    except Exception as e:
                        zf.writestr(f"ERROR_FILA_{idx}.txt", f"Error en fila {idx}: {e}")
            zip_buffer.seek(0)
            st.download_button(
                "⬇️ Descargar ZIP con todos los documentos",
                data=zip_buffer,
                file_name=f"Documentos_{subetapa}.zip",
                mime="application/zip"
            )
# ============================================================
# 🧩 TAB 2 — Generador de Demandas y Medidas Cautelares
# ============================================================
with tab2:
    st.header("📄 Generador de Demandas y Medidas Cautelares")

    # OJO: el nombre tiene un espacio en el archivo original del usuario
    TEMPLATE_DEMANDA_PATH = "FORMATO_DEMANDA_FINAL.docx"
    TEMPLATE_MEDIDAS_PATH = "FORMATO_ SOLICITUD_MEDIDAS_FINAL.docx"

    REQUIRED_COLUMNS = [
        "CC", "NOMBRE", "VALIDACION", "PAGARE",
        "FECHA_VENCIMIENTO", "FECHA_INTERESES",
        "JUZGADO", "CUANTIA",
        "CAPITAL", "CAPITAL_EN_LETRAS",
        "DOMICILIO", "CIUDAD",
    ]

    def make_excel_template_bytes():
        df = pd.DataFrame(columns=REQUIRED_COLUMNS)
        bio = io.BytesIO()
        with pd.ExcelWriter(bio, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="BASE")
        bio.seek(0)
        return bio

    def sanitize_filename(s: str) -> str:
        s = str(s or "").strip()
        s = re.sub(r"\s+", "_", s)
        s = re.sub(r"[^A-Za-z0-9_\-\.]", "", s)
        return s

    def juzgado_con_reparto(juzgado_text: str) -> str:
        jt = str(juzgado_text or "").strip()
        jt_clean = re.sub(r"\(REPARTO\)", "", jt, flags=re.IGNORECASE).strip()
        return f"{jt_clean}\n(REPARTO)"

    def replace_placeholders_doc(doc: Document, mapping: dict):
        for p in doc.paragraphs:
            txt = p.text
            for k, v in mapping.items():
                txt = txt.replace(f"{{{k}}}", str(v))
            p.text = txt
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = p.text
                        for k, v in mapping.items():
                            txt = txt.replace(f"{{{k}}}", str(v))
                        p.text = txt

    # ---------- UI ----------
    c1, c2 = st.columns([2, 1])
    with c1:
        excel_file = st.file_uploader("📂 Cargar base Excel (.xlsx / .xlsm)", type=["xlsx", "xlsm"])
    with c2:
        st.download_button(
            "📘 Descargar archivo guía",
            data=make_excel_template_bytes(),
            file_name="PLANTILLA_CRUCE_DEMANDAS_MMCC_OFICIAL.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    if excel_file:
        try:
            df = pd.read_excel(excel_file)
        except Exception as e:
            st.error(f"No se pudo leer el Excel: {e}")
            st.stop()

        faltantes = [c for c in REQUIRED_COLUMNS if c not in list(df.columns)]
        if faltantes:
            st.error(f"Faltan columnas: {faltantes}")
        else:
            # 💰 Formato moneda (solo valor, sin 'COP')
            try:
                df["CAPITAL"] = (
                    df["CAPITAL"]
                    .astype(str)
                    .replace({r"[\$,]": "", r"\.": ""}, regex=True)
                    .replace("", "0")
                    .astype(float)
                )
                df["CAPITAL"] = df["CAPITAL"].apply(lambda x: f"${x:,.0f}".replace(",", "."))
            except Exception as e:
                st.warning(f"No se pudo formatear CAPITAL: {e}")

            st.success(f"Base cargada: {df.shape[0]} registros")
            st.dataframe(df.head(3))

            tipo_doc = st.selectbox("📄 Modelo a generar", ["DEMANDA", "MEDIDAS"])
            indices = list(df.index)
            sel_idx = st.selectbox(
                "🔎 Registro individual",
                indices,
                format_func=lambda i: f"{i} — {df.loc[i, 'NOMBRE']} — CC {df.loc[i, 'CC']}"
            )

            row = df.loc[sel_idx]

            mapping_preview = {
                "JUZGADO": juzgado_con_reparto(row.get("JUZGADO", "")),
                "CUANTIA": row.get("CUANTIA", ""),
                "NOMBRE": row.get("NOMBRE", ""),
                "CC": row.get("CC", ""),
                "CIUDAD": row.get("CIUDAD", ""),
                "PAGARE": row.get("PAGARE", ""),
                "CAPITAL_EN_LETRAS": row.get("CAPITAL_EN_LETRAS", ""),
                "CAPITAL": row.get("CAPITAL", ""),
                "FECHA_VENCIMIENTO": row.get("FECHA_VENCIMIENTO", ""),
                "FECHA_INTERESES": row.get("FECHA_INTERESES", ""),
                "DOMICILIO": row.get("DOMICILIO", ""),
            }

            st.markdown("#### 👀 Vista previa")
            st.code(
f"""Señor:
{mapping_preview.get("JUZGADO", "").splitlines()[0]}
(REPARTO)
E. S. D.

REFERENCIA: PROCESO EJECUTIVO DE {mapping_preview.get("CUANTIA","")} CUANTÍA.
DEMANDANTE : BANCO GNB SUDAMERIS S.A.
DEMANDADO : {mapping_preview.get("NOMBRE","")} CC {mapping_preview.get("CC","")}
""", language="markdown")

            # ---------- Generar individual ----------
            if st.button("📄 Generar documento individual"):
                tpl_path = TEMPLATE_DEMANDA_PATH if tipo_doc == "DEMANDA" else TEMPLATE_MEDIDAS_PATH
                doc = Document(tpl_path)
                replace_placeholders_doc(doc, mapping_preview)
                out_name = f"{sanitize_filename(mapping_preview['CC'])}_{sanitize_filename(mapping_preview['NOMBRE'])}_{tipo_doc}.docx"
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                st.download_button("⬇️ Descargar documento", data=bio,
                                   file_name=out_name,
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            st.divider()
            st.markdown("### 📦 Generación masiva")
            st.write("Genera todos los documentos en un solo archivo ZIP.")

            if st.button("📦 Generar TODOS (ZIP)"):
                tpl_path = TEMPLATE_DEMANDA_PATH if tipo_doc == "DEMANDA" else TEMPLATE_MEDIDAS_PATH
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    for idx, fila in df.iterrows():
                        mapping = {
                            "JUZGADO": juzgado_con_reparto(fila.get("JUZGADO", "")),
                            "CUANTIA": fila.get("CUANTIA", ""),
                            "NOMBRE": fila.get("NOMBRE", ""),
                            "CC": fila.get("CC", ""),
                            "CIUDAD": fila.get("CIUDAD", ""),
                            "PAGARE": fila.get("PAGARE", ""),
                            "CAPITAL_EN_LETRAS": fila.get("CAPITAL_EN_LETRAS", ""),
                            "CAPITAL": fila.get("CAPITAL", ""),
                            "FECHA_VENCIMIENTO": fila.get("FECHA_VENCIMIENTO", ""),
                            "FECHA_INTERESES": fila.get("FECHA_INTERESES", ""),
                            "DOMICILIO": fila.get("DOMICILIO", ""),
                        }
                        try:
                            doc = Document(tpl_path)
                            replace_placeholders_doc(doc, mapping)
                            out_name = f"{sanitize_filename(mapping['CC'])}_{sanitize_filename(mapping['NOMBRE'])}_{tipo_doc}.docx"
                            mem = io.BytesIO()
                            doc.save(mem)
                            mem.seek(0)
                            zf.writestr(out_name, mem.read())
                        except Exception as e:
                            zf.writestr(f"ERROR_FILA_{idx}.txt", f"Error en fila {idx}: {e}")
                zip_buffer.seek(0)
                st.download_button(
                    "⬇️ Descargar ZIP con todos los documentos",
                    data=zip_buffer,
                    file_name=f"Documentos_{tipo_doc}.zip",
                    mime="application/zip"
                )
    else:
        st.info("Sube el Excel para continuar.")
